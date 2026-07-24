"""
test_pipeline.py — Tests the updated ResumeTailor pipeline.
"""

import json
import sys
import io
import re
from pathlib import Path

# Insert current directory to path
sys.path.insert(0, str(Path(__file__).parent))

# ── Mock tailored resume (what Gemini would return) ──────────────────────
MOCK_GEMINI_OUTPUT = {
    "tailored_resume": {
        "summary": (
            "Software Engineer with hands-on experience designing and shipping scalable "
            "backend services, REST APIs, and microservices using Python, Java, and Go."
        ),
        "skills_selected": [
            {"category": "Languages", "items": ["Python", "JavaScript", "Go", "SQL", "Bash"]},
            {"category": "Frameworks", "items": ["FastAPI", "Django", "React", "Pytest"]},
            {"category": "Databases", "items": ["PostgreSQL", "MongoDB", "Redis"]},
            {"category": "Cloud & DevOps", "items": ["Docker", "Kubernetes", "GitHub Actions"]},
            {"category": "Tools", "items": ["Git", "REST APIs", "Postman"]},
        ],
        "projects": [
            {
                "name": "SecureVault – Encrypted Secret Manager",
                "section": "Projects",
                "dates": "Jan 2024 – Mar 2024",
                "tech_stack": "Python, FastAPI, PostgreSQL, Redis, Docker, JWT, AES-256",
                "github_link": "https://github.com/alexjordan/securevault",
                "bullets": [
                    "Designed RESTful FastAPI service with JWT-authenticated endpoints, rate limiting via Redis.",
                    "Containerized the full stack with Docker Compose; wrote GitHub Actions CI/CD pipeline.",
                    "Built a self-hosted secrets manager with AES-256 encryption, role-based access control.",
                ],
            },
            {
                "name": "AutoTest – AI-Assisted QA Framework",
                "section": "Projects",
                "dates": "May 2023 – Aug 2023",
                "tech_stack": "Python, Playwright, Selenium, Pytest, Jenkins, GitHub Actions",
                "github_link": "https://github.com/alexjordan/autotest",
                "bullets": [
                    "Developed an end-to-end test automation framework using Playwright and Pytest.",
                    "Configured Jenkins pipelines and GitHub Actions workflows to run the Playwright suite.",
                ],
            },
        ],
        "internships_selected": [
            {
                "company": "TechCorp Inc.",
                "role": "Software Engineering Intern",
                "dates": "Jun 2022 – Aug 2022",
                "bullets": [
                    "Developed and shipped 3 new REST API endpoints in core Java Spring Boot.",
                    "Optimized a slow PostgreSQL query that ran on the nightly batch job.",
                    "Wrote unit and integration tests with JUnit and Mockito.",
                ],
            }
        ],
        "education": [
            {
                "institution": "University of California, Berkeley",
                "degree": "Bachelor of Science in Computer Science",
                "dates": "Aug 2019 – May 2023",
                "gpa": "3.8/4.0",
                "relevant_coursework": "Data Structures, Algorithms, Operating Systems, Computer Networks",
            }
        ],
        "certifications": [
            {"name": "AWS Certified Developer – Associate", "issuer": "Amazon Web Services", "date": "March 2024"},
            {"name": "Google Associate Cloud Engineer", "issuer": "Google Cloud", "date": "November 2023"},
        ],
        "publications": [],
        "optional_sections": [],
        "section_order": ["summary", "skills", "projects", "internships", "education", "certifications"],
        "section_titles": {
            "summary": "Professional Summary",
            "skills": "Technical Skills",
            "projects": "Projects",
            "internships": "Work Experience",
            "education": "Education",
            "certifications": "Certifications"
        },
        "contact": {
            "name": "Alex Jordan",
            "email": "alex.jordan@email.com",
            "phone": "+1 (555) 123-4567",
            "location": "San Francisco, CA"
        }
    },
    "entry_counts": {
        "master_projects_count": 2,
        "master_internships_count": 1,
        "master_education_count": 1,
        "master_certifications_count": 2,
        "master_publications_count": 0
    },
    "dashboard": {
        "ats_score": 90,
        "ats_explanation": "Great formatting.",
        "readability_score": 90,
        "readability_explanation": "Very readable.",
        "match_score": 85,
        "match_explanation": "Good fit.",
        "keyword_coverage": 80,
        "missing_skills": [],
        "weaknesses": [],
        "strengths": ["FastAPI", "PostgreSQL"],
        "improvements": []
    }
}

# ── Dynamic mock master resume text ──────────────────────────────────────
MOCK_MASTER_TEXT = """
Alex Jordan
alex.jordan@email.com | +1 (555) 123-4567 | San Francisco, CA
GitHub: https://github.com/alexjordan | Portfolio: https://alexjordan.dev | LinkedIn: https://linkedin.com/in/alexjordan

SUMMARY:
Software Engineer with hands-on experience designing and shipping scalable backend services, REST APIs, and microservices using Python, Java, and Go.

TECHNICAL SKILLS:
Languages: Python, JavaScript, TypeScript, Java, Go, SQL, Bash
Frameworks: FastAPI, Django, React, Pytest
Databases: PostgreSQL, MongoDB, Redis
Cloud & DevOps: Docker, Kubernetes, GitHub Actions
Tools: Git, REST APIs, Postman

WORK EXPERIENCE:
TechCorp Inc. - Software Engineering Intern (Jun 2022 – Aug 2022)
* Developed and shipped 3 new REST API endpoints in core Java Spring Boot.
* Optimized a slow PostgreSQL query.
* Wrote unit and integration tests with JUnit and Mockito.

PROJECTS:
SecureVault – Encrypted Secret Manager (Jan 2024 – Mar 2024)
Tech: Python, FastAPI, PostgreSQL, Redis, Docker, JWT, AES-256
* Built a self-hosted secrets manager.
* Designed RESTful FastAPI service.
* Containerized the full stack.

AutoTest – AI-Assisted QA Framework (May 2023 – Aug 2023)
Tech: Python, Playwright, Selenium, Pytest, Jenkins, GitHub Actions
* Developed automation framework.
* Configured Jenkins pipelines.

EDUCATION:
University of California, Berkeley - Bachelor of Science in Computer Science (Aug 2019 – May 2023)

CERTIFICATIONS:
AWS Certified Developer – Associate - Amazon Web Services (March 2024)
Google Associate Cloud Engineer - Google Cloud (November 2023)
"""

def make_dummy_pdf() -> bytes:
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Alex Jordan")
    c.drawString(100, 730, "alex.jordan@email.com | +1 (555) 123-4567")
    c.drawString(100, 710, "Python, FastAPI, PostgreSQL, Redis, Docker, CI/CD")
    c.save()
    return buf.getvalue()

def make_dummy_docx() -> bytes:
    import docx
    doc = docx.Document()
    doc.add_paragraph("Alex Jordan")
    doc.add_paragraph("alex.jordan@email.com | +1 (555) 123-4567")
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Redis, Docker, CI/CD")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Test 1: File text extraction ──────────────────────────────────────────
print("\n" + "="*60)
print("TEST 1: File text extraction (_extract_text_from_file)")
print("="*60)

from main import _extract_text_from_file

# Test PDF Extraction
pdf_bytes = make_dummy_pdf()
extracted_pdf_text = _extract_text_from_file("my_resume.pdf", pdf_bytes)
print(f"[OK] Extracted PDF text length: {len(extracted_pdf_text)} characters")
assert "Alex Jordan" in extracted_pdf_text, "Failed to extract name from PDF"

# Test DOCX Extraction
docx_bytes = make_dummy_docx()
extracted_docx_text = _extract_text_from_file("my_resume.docx", docx_bytes)
print(f"[OK] Extracted DOCX text length: {len(extracted_docx_text)} characters")
assert "Alex Jordan" in extracted_docx_text, "Failed to extract name from DOCX"

# Test Scanned PDF/Empty PDF error handling
try:
    _extract_text_from_file("empty.pdf", b"%PDF-1.4 ... empty bytes ...")
    assert False, "Should have raised a ValueError for bad/scanned PDF"
except ValueError as exc:
    print(f"[OK] Caught expected error for scanned/bad PDF: '{exc}'")
    assert "please make sure it's a text-based PDF" in str(exc)


# ── Test 2: Fabrication & Count validation check ────────────────────────
print("\n" + "="*60)
print("TEST 2: Fabrication & Count validation check")
print("="*60)

from gemini_service import validate_tailored_resume

# Should produce ZERO errors
errors = validate_tailored_resume(MOCK_GEMINI_OUTPUT, MOCK_MASTER_TEXT)
print(f"[OK] Clean response: {len(errors)} validation errors (expected 0)")
assert len(errors) == 0, f"Unexpected errors: {errors}"

# Test count check failure
MOCK_WITH_BAD_COUNT = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_WITH_BAD_COUNT["entry_counts"]["master_projects_count"] = 5
errors_with_bad_count = validate_tailored_resume(MOCK_WITH_BAD_COUNT, MOCK_MASTER_TEXT)
print(f"[OK] Mismatched projects count -> error caught: {errors_with_bad_count}")
assert len(errors_with_bad_count) > 0, "Failed to catch projects count mismatch"

# Test fabrication check failure
MOCK_WITH_FAKE = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_WITH_FAKE["tailored_resume"]["skills_selected"].append({
    "category": "Fake Skills",
    "items": ["QuantumNeuralSynapse"],
})
errors_with_fake, unverified_with_fake = validate_tailored_resume(MOCK_WITH_FAKE, MOCK_MASTER_TEXT, return_unverified=True)
print(f"[OK] Injected fake skill -> warning caught: {unverified_with_fake}")
assert len(unverified_with_fake) > 0, "Failed to catch fabricated skill"
assert unverified_with_fake[0]["item"] == "QuantumNeuralSynapse"

# Test summary buzzword check failure
MOCK_WITH_BUZZWORD = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_WITH_BUZZWORD["tailored_resume"]["summary"] = "A highly motivated candidate with a proven track record."
errors_buzz = validate_tailored_resume(MOCK_WITH_BUZZWORD, MOCK_MASTER_TEXT)
print(f"[OK] Buzzword summary -> error caught: {errors_buzz}")
assert len(errors_buzz) > 0, "Failed to catch summary buzzword phrase"

# Test skills floor failure (empty)
MOCK_WITH_EMPTY_SKILLS = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_WITH_EMPTY_SKILLS["tailored_resume"]["skills_selected"] = []
errors_empty_skills = validate_tailored_resume(MOCK_WITH_EMPTY_SKILLS, MOCK_MASTER_TEXT)
print(f"[OK] Empty skills -> error caught: {errors_empty_skills}")
assert len(errors_empty_skills) > 0, "Failed to catch empty skills list"

# Test skills floor failure (sparse)
MOCK_WITH_SPARSE_SKILLS = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_WITH_SPARSE_SKILLS["tailored_resume"]["skills_selected"] = [
    {"category": "Tech", "items": ["Python", "Git"]}
]
errors_sparse_skills = validate_tailored_resume(MOCK_WITH_SPARSE_SKILLS, MOCK_MASTER_TEXT)
print(f"[OK] Sparse skills (< 5 items) -> error caught: {errors_sparse_skills}")
assert len(errors_sparse_skills) > 0, "Failed to catch sparse skills list"

# Test certifications independent count mismatch
MOCK_MASTER_WITH_CERTS = """
CERTIFICATIONS:
AWS Certified Developer – Associate
Google Associate Cloud Engineer
Certified Kubernetes Administrator
Offensive Security Certified Professional
CompTIA Security+
"""
MOCK_TAILORED_WITH_MERGED_CERTS = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_TAILORED_WITH_MERGED_CERTS["tailored_resume"]["certifications"] = [
    {"name": "All 5 certifications merged", "issuer": "Various", "date": "2024"}
]
errors_merged_certs = validate_tailored_resume(MOCK_TAILORED_WITH_MERGED_CERTS, MOCK_MASTER_WITH_CERTS)
print(f"[OK] Merged certifications -> error caught: {errors_merged_certs}")
assert len(errors_merged_certs) > 0, "Failed to catch merged certifications"
assert "Certifications section: master resume has approximately 5 entries" in errors_merged_certs[0]

# Test 2.5: Project links preservation check
MOCK_MASTER_WITH_PROJ_LINKS = """
PROJECTS:
SecureVault – Encrypted Secret Manager (Jan 2024 – Mar 2024)
* GitHub: https://github.com/alexjordan/securevault
* Live Demo: https://securevault.live
"""
MOCK_TAILORED_WITH_NO_LINKS = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
for proj in MOCK_TAILORED_WITH_NO_LINKS["tailored_resume"]["projects"]:
    proj["github_link"] = ""
    proj["project_url"] = ""
    proj["live_demo"] = ""
    proj["links"] = []
    proj["additional_links"] = []

errors_missing_links = validate_tailored_resume(MOCK_TAILORED_WITH_NO_LINKS, MOCK_MASTER_WITH_PROJ_LINKS)
print(f"[OK] Missing project links -> auto-restored without errors: {errors_missing_links}")
assert len(errors_missing_links) == 0, f"Expected links to be auto-restored without errors, got: {errors_missing_links}"
p = MOCK_TAILORED_WITH_NO_LINKS["tailored_resume"]["projects"][0]
assert p.get("github_link") == "https://github.com/alexjordan/securevault", f"Expected github_link to be restored, got: {p.get('github_link')}"
assert p.get("live_demo") == "https://securevault.live", f"Expected live_demo to be restored, got: {p.get('live_demo')}"

# Test 2.6: Section order normalization check
mock_raw_json = {
    "tailored_resume": {
        "summary": "Experienced engineer.",
        "skills_selected": [{"category": "Languages", "items": ["Python", "Java", "SQL", "Go", "C++"]}],
        "projects": [],
        "internships_selected": [],
        "education": [],
        "certifications": [],
        "section_order": ["skills_selected", "internships_selected"],
        "section_titles": {
            "summary": "Summary",
            "skills": "Skills",
            "projects": "Projects",
            "internships": "Experience",
            "education": "Education",
            "certifications": "Certifications",
            "publications": "Publications"
        }
    },
    "entry_counts": {
        "master_projects_count": 0,
        "master_internships_count": 0,
        "master_education_count": 0,
        "master_certifications_count": 0,
        "master_publications_count": 0
    },
    "dashboard": {}
}

raw_order = mock_raw_json["tailored_resume"]["section_order"]
normalized_order = []
for sec in raw_order:
    sec_lower = sec.lower().strip()
    if sec_lower in ("skills_selected", "skills", "technical skills"):
        normalized_order.append("skills")
    elif sec_lower in ("internships_selected", "internships", "experience", "work experience", "professional experience"):
        normalized_order.append("internships")
    else:
        normalized_order.append(sec_lower)

print(f"[OK] Normalized section order: {normalized_order}")
assert normalized_order == ["skills", "internships"], f"Normalization failed: {normalized_order}"


# Test 2.7: Fabrication check - singular/plural tolerance (positive test)
MOCK_MASTER_REST_API = "Skills:\nREST API design\nAWS Certified Developer"
MOCK_TAILORED_REST_APIS = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_TAILORED_REST_APIS["tailored_resume"]["projects"] = []
MOCK_TAILORED_REST_APIS["entry_counts"]["master_projects_count"] = 0
MOCK_TAILORED_REST_APIS["tailored_resume"]["skills_selected"] = [
    {"category": "Tools & Design", "items": ["REST APIs"]}
]
errors_plural_match = validate_tailored_resume(MOCK_TAILORED_REST_APIS, MOCK_MASTER_REST_API)
fabrication_errors = [e for e in errors_plural_match if "Fabrication check failed" in e]
print(f"[OK] Fabrication check with plural/singular tolerance -> errors: {fabrication_errors}")
assert len(fabrication_errors) == 0, f"Expected no fabrication error for REST APIs, got: {fabrication_errors}"

# Test 2.8: Fabrication check - genuine fabrication rejection (negative test)
MOCK_MASTER_NO_KUBERNETES = "Skills:\nREST API design"
MOCK_TAILORED_WITH_KUBERNETES = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_TAILORED_WITH_KUBERNETES["tailored_resume"]["projects"] = []
MOCK_TAILORED_WITH_KUBERNETES["entry_counts"]["master_projects_count"] = 0
MOCK_TAILORED_WITH_KUBERNETES["tailored_resume"]["skills_selected"] = [
    {"category": "Tools & Design", "items": ["Kubernetes"]}
]
errors_kubernetes, unverified_kubernetes = validate_tailored_resume(MOCK_TAILORED_WITH_KUBERNETES, MOCK_MASTER_NO_KUBERNETES, return_unverified=True)
print(f"[OK] Genuine fabrication check for Kubernetes -> unverified: {unverified_kubernetes}")
assert len(unverified_kubernetes) > 0, "Expected fabrication warning for Kubernetes, but passed!"
assert unverified_kubernetes[0]["item"] == "Kubernetes"

# Test 2.9: Fabrication check - user-selected keyword exemption (positive test)
MOCK_MASTER_NO_OOP = "Skills:\nREST API design"
MOCK_TAILORED_WITH_OOP = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_TAILORED_WITH_OOP["tailored_resume"]["projects"] = []
MOCK_TAILORED_WITH_OOP["entry_counts"]["master_projects_count"] = 0
MOCK_TAILORED_WITH_OOP["tailored_resume"]["skills_selected"] = [
    {"category": "Tools & Design", "items": ["Object-oriented programming"]}
]
errors_oop = validate_tailored_resume(MOCK_TAILORED_WITH_OOP, MOCK_MASTER_NO_OOP, ["Object-oriented programming"])
fabrication_errors_oop = [e for e in errors_oop if "Fabrication check failed" in e]
print(f"[OK] Fabrication check with user keyword exemption -> errors: {fabrication_errors_oop}")
assert len(fabrication_errors_oop) == 0, f"Expected no fabrication error for Object-oriented programming, got: {fabrication_errors_oop}"

# Test 2.10: Fabrication check - genuine fabrication with unselected keyword (negative test)
MOCK_MASTER_NO_KUBERNETES_SEL = "Skills:\nREST API design"
MOCK_TAILORED_WITH_KUBERNETES_SEL = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_TAILORED_WITH_KUBERNETES_SEL["tailored_resume"]["projects"] = []
MOCK_TAILORED_WITH_KUBERNETES_SEL["entry_counts"]["master_projects_count"] = 0
MOCK_TAILORED_WITH_KUBERNETES_SEL["tailored_resume"]["skills_selected"] = [
    {"category": "Tools & Design", "items": ["Kubernetes"]}
]
errors_kubernetes_sel, unverified_kubernetes_sel = validate_tailored_resume(MOCK_TAILORED_WITH_KUBERNETES_SEL, MOCK_MASTER_NO_KUBERNETES_SEL, ["REST API"], return_unverified=True)
print(f"[OK] Fabrication check with different user keyword -> unverified: {unverified_kubernetes_sel}")
assert len(unverified_kubernetes_sel) > 0, "Expected fabrication warning for Kubernetes, but passed!"
assert unverified_kubernetes_sel[0]["item"] == "Kubernetes"


# Test 2.11: Project title verification and correction test
MOCK_MASTER_PROJECT_TITLE = """
PROJECTS:
1. SecureVault — Encrypted Secret Manager Jan 2024 – Mar 2024
"""
MOCK_TAILORED_SHORT_TITLE = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_TAILORED_SHORT_TITLE["tailored_resume"]["projects"] = [
    {
        "name": "SecureVault",
        "tech_stack": "Python, FastAPI",
        "bullets": ["Wrote some code."]
    }
]
MOCK_TAILORED_SHORT_TITLE["entry_counts"]["master_projects_count"] = 1

errors_title = validate_tailored_resume(MOCK_TAILORED_SHORT_TITLE, MOCK_MASTER_PROJECT_TITLE)
corrected_title = MOCK_TAILORED_SHORT_TITLE["tailored_resume"]["projects"][0]["name"]
print(f"[OK] Project title automatically corrected: '{corrected_title}'")
assert corrected_title == "SecureVault — Encrypted Secret Manager", f"Project title was not corrected: '{corrected_title}'"


# Test 2.12: Selected keywords missing verification test
MOCK_TAILORED_NO_KEYWORDS = json.loads(json.dumps(MOCK_GEMINI_OUTPUT))
MOCK_TAILORED_NO_KEYWORDS["tailored_resume"]["summary"] = "Experienced backend developer."
MOCK_TAILORED_NO_KEYWORDS["tailored_resume"]["skills_selected"] = []
errors_missing_kws = validate_tailored_resume(
    MOCK_TAILORED_NO_KEYWORDS,
    MOCK_MASTER_TEXT,
    ["Cloud technologies", "Data structures"]
)
print(f"[OK] Missing selected keywords -> error caught: {errors_missing_kws}")
assert any("Missing user-selected keywords" in err for err in errors_missing_kws), "Failed to catch missing selected keywords"
assert "Cloud technologies" in "".join(errors_missing_kws)
assert "Data structures" in "".join(errors_missing_kws)


# Test 2.13: Simulating repeated buzzword failures across 3 attempts before succeeding on the 4th
print("\nTEST 2.13: Buzzword retry budget (4 attempts) simulation")
import gemini_service
original_generate_content = gemini_service.generate_content_with_fallback

class MockResponse:
    def __init__(self, text):
        self.text = text

attempt_count = 0
outputs_buzzword = [
    # Attempt 1
    {
        "tailored_resume": {
            "summary": "Proven ability to write software.",
            "skills_selected": [{"category": "L", "items": ["Python", "Java", "SQL", "C", "C++"]}],
            "projects": [],
            "internships_selected": [],
            "education": [],
            "certifications": [],
            "section_order": ["skills"],
            "section_titles": {}
        },
        "entry_counts": {
            "master_projects_count": 0,
            "master_internships_count": 0,
            "master_education_count": 0
        },
        "dashboard": {}
    },
    # Attempt 2
    {
        "tailored_resume": {
            "summary": "Proven track record of success.",
            "skills_selected": [{"category": "L", "items": ["Python", "Java", "SQL", "C", "C++"]}],
            "projects": [],
            "internships_selected": [],
            "education": [],
            "certifications": [],
            "section_order": ["skills"],
            "section_titles": {}
        },
        "entry_counts": {
            "master_projects_count": 0,
            "master_internships_count": 0,
            "master_education_count": 0
        },
        "dashboard": {}
    },
    # Attempt 3
    {
        "tailored_resume": {
            "summary": "Highly motivated engineer.",
            "skills_selected": [{"category": "L", "items": ["Python", "Java", "SQL", "C", "C++"]}],
            "projects": [],
            "internships_selected": [],
            "education": [],
            "certifications": [],
            "section_order": ["skills"],
            "section_titles": {}
        },
        "entry_counts": {
            "master_projects_count": 0,
            "master_internships_count": 0,
            "master_education_count": 0
        },
        "dashboard": {}
    },
    # Attempt 4
    {
        "tailored_resume": {
            "summary": "Experienced engineer with database skills.",
            "skills_selected": [{"category": "L", "items": ["Python", "Java", "SQL", "C", "C++"]}],
            "projects": [],
            "internships_selected": [],
            "education": [],
            "certifications": [],
            "section_order": ["skills"],
            "section_titles": {}
        },
        "entry_counts": {
            "master_projects_count": 0,
            "master_internships_count": 0,
            "master_education_count": 0
        },
        "dashboard": {}
    }
]

def mock_generate_buzzword(*args, **kwargs):
    global attempt_count
    idx = min(attempt_count, len(outputs_buzzword) - 1)
    attempt_count += 1
    return MockResponse(json.dumps(outputs_buzzword[idx]))

gemini_service.generate_content_with_fallback = mock_generate_buzzword
import os
original_api_key = os.environ.get("GEMINI_API_KEY")
os.environ["GEMINI_API_KEY"] = "mock_key"

try:
    res = gemini_service.generate_tailored_resume("Skills:\nPython\nJava\nSQL\nC\nC++", "Job requiring Python")
    print(f"[OK] Buzzword retry simulation succeeded on attempt {attempt_count}! Output summary: '{res['tailored_resume']['summary']}'")
    assert attempt_count == 4, f"Expected 4 attempts, but did: {attempt_count}"
    assert "Proven" not in res['tailored_resume']['summary'], "Summary still has buzzword"
except Exception as e:
    print(f"Error: {e}")
    assert False, f"Buzzword retry simulation failed: {e}"
finally:
    gemini_service.generate_content_with_fallback = original_generate_content
    if original_api_key is not None:
        os.environ["GEMINI_API_KEY"] = original_api_key


# Test 2.14: Simulating non-buzzword failures (project count mismatch) respecting the 2-attempt limit
print("\nTEST 2.14: Non-buzzword failure respects 2-attempt budget")
attempt_count_fab = 0
outputs_fabrication = [
    # Attempt 1: contains project count mismatch (fail)
    {
        "tailored_resume": {
            "summary": "Experienced engineer.",
            "skills_selected": [{"category": "L", "items": ["Python", "Java", "SQL", "C", "C++"]}],
            "projects": [],
            "internships_selected": [],
            "education": [],
            "certifications": [],
            "section_order": ["skills"],
            "section_titles": {}
        },
        "entry_counts": {
            "master_projects_count": 2,
            "master_internships_count": 0,
            "master_education_count": 0
        },
        "dashboard": {}
    },
    # Attempt 2: still contains project count mismatch (fail)
    {
        "tailored_resume": {
            "summary": "Experienced engineer.",
            "skills_selected": [{"category": "L", "items": ["Python", "Java", "SQL", "C", "C++"]}],
            "projects": [],
            "internships_selected": [],
            "education": [],
            "certifications": [],
            "section_order": ["skills"],
            "section_titles": {}
        },
        "entry_counts": {
            "master_projects_count": 2,
            "master_internships_count": 0,
            "master_education_count": 0
        },
        "dashboard": {}
    },
    # Attempt 3: (should not be reached)
    {
        "tailored_resume": {
            "summary": "Clean developer.",
            "skills_selected": [{"category": "L", "items": ["Python", "Java", "SQL", "C", "C++"]}],
            "projects": [],
            "internships_selected": [],
            "education": [],
            "certifications": [],
            "section_order": ["skills"],
            "section_titles": {}
        },
        "entry_counts": {
            "master_projects_count": 0,
            "master_internships_count": 0,
            "master_education_count": 0
        },
        "dashboard": {}
    }
]

def mock_generate_fabrication(*args, **kwargs):
    global attempt_count_fab
    idx = min(attempt_count_fab, len(outputs_fabrication) - 1)
    attempt_count_fab += 1
    return MockResponse(json.dumps(outputs_fabrication[idx]))

gemini_service.generate_content_with_fallback = mock_generate_fabrication

try:
    res = gemini_service.generate_tailored_resume("Skills:\nPython\nJava\nSQL\nC\nC++\nProjects:\n- Proj 1\n- Proj 2", "Job requiring Python")
    assert False, "Expected count mismatch retry check to fail after 2 attempts, but it succeeded!"
except ValueError as e:
    print(f"[OK] Non-buzzword retry stopped correctly at 2 attempts! Attempt count: {attempt_count_fab}")
    assert attempt_count_fab == 2, f"Expected exactly 2 attempts, but did: {attempt_count_fab}"
    assert "Entry count mismatch" in str(e), f"Expected count mismatch error, got: {e}"
finally:
    gemini_service.generate_content_with_fallback = original_generate_content


# Test 2.15: Deduplication of near-identical skill items
print("\nTEST 2.15: Deduplication of near-identical skill items")
from gemini_service import deduplicate_skills_selected
mock_skills_dup = [
    {
        "category": "Tools",
        "items": ["REST API", "Docker", "REST APIs", "Docker ", "databases", "Database"]
    }
]
deduped_skills = deduplicate_skills_selected(mock_skills_dup)
items = deduped_skills[0]["items"]
print(f"[OK] Deduped items: {items}")
assert len(items) == 3, f"Expected exactly 3 items, got: {items}"
assert "REST API" in items, "Expected 'REST API' to be kept (first encountered)"
assert "REST APIs" not in items, "Expected 'REST APIs' to be removed as near-duplicate"
assert "Docker" in items
assert "databases" in items
assert "Database" not in items

# Test 2.16: Confirm genuinely different skills are NOT incorrectly merged
print("\nTEST 2.16: Non-duplicates are not merged")
mock_skills_diff = [
    {
        "category": "Tools",
        "items": ["REST API", "GraphQL API", "APIs", "Database Management"]
    }
]
deduped_diff = deduplicate_skills_selected(mock_skills_diff)
items_diff = deduped_diff[0]["items"]
print(f"[OK] Non-deduped items: {items_diff}")
assert len(items_diff) == 4, f"Expected 4 items, got: {items_diff}"
assert "GraphQL API" in items_diff
assert "REST API" in items_diff


# Test 2.19: Project link restoration
print("\nTEST 2.19: Project link restoration")
from gemini_service import restore_project_links

mock_master_links = """
PROJECTS:
ResumeTailor AI — AI-Powered, ATS-Safe Resume Customization Engine Jun 2026 – Jul 2026
* GitHub: github.com/Iamgokul7/ResumeTailor-AI | Live: resume-tailorj5zl.onrender.com
"""

mock_tailored_no_live = {
    "tailored_resume": {
        "projects": [
            {
                "name": "ResumeTailor AI — AI-Powered, ATS-Safe Resume Customization Engine",
                "github_link": "github.com/Iamgokul7/ResumeTailor-AI",
                "live_demo": ""
            }
        ]
    }
}

restore_project_links(mock_tailored_no_live, mock_master_links)
restored_live = mock_tailored_no_live["tailored_resume"]["projects"][0].get("live_demo")
print(f"[OK] Restored live demo link: {restored_live}")
assert restored_live == "resume-tailorj5zl.onrender.com", f"Expected live demo to be restored, got: {restored_live}"


# Test 2.20: Cross-category deduplication
print("\nTEST 2.20: Cross-category deduplication")
mock_skills_cross = [
    {
        "category": "Web Development",
        "items": ["HTML", "CSS", "JavaScript"]
    },
    {
        "category": "Tools & Design",
        "items": ["HTML", "CSS", "Git", "C++"]
    },
    {
        "category": "Languages",
        "items": ["Python", "Java", "C++"]
    }
]

deduped_cross = deduplicate_skills_selected(mock_skills_cross)
cat1 = deduped_cross[0]["items"]
cat2 = deduped_cross[1]["items"]
cat3 = deduped_cross[2]["items"]

print(f"Cat 1: {cat1}")
print(f"Cat 2: {cat2}")
print(f"Cat 3: {cat3}")

assert "HTML" in cat1
assert "CSS" in cat1
assert "HTML" not in cat2
assert "CSS" not in cat2
assert "C++" in cat3
assert "C++" not in cat2
assert "JavaScript" in cat1
print("[OK] Cross-category deduplication works perfectly!")


# Test 2.21: Classification-based skill filtering
print("\nTEST 2.21: Classification-based skill filtering")
from gemini_service import filter_non_skills_via_classification

class MockGenaiModels:
    def generate_content(self, model, contents, config):
        data = [
            {"item": "Wireshark", "classification": "TECHNICAL_SKILL"},
            {"item": "DWT Steganography", "classification": "TECHNICAL_SKILL"},
            {"item": "JWT Authentication", "classification": "TECHNICAL_SKILL"},
            {"item": "Stakeholder Communication", "classification": "NOT_A_SKILL"},
            {"item": "Requirements Gathering", "classification": "NOT_A_SKILL"}
        ]
        class MockResponseText:
            def __init__(self, text):
                self.text = text
        return MockResponseText(json.dumps(data))

class MockClient:
    def __init__(self):
        self.models = MockGenaiModels()

mock_tailored = {
    "tailored_resume": {
        "skills_selected": [
            {
                "category": "Skills",
                "items": ["Wireshark", "DWT Steganography", "JWT Authentication", "Stakeholder Communication", "Requirements Gathering"]
            }
        ]
    }
}

filter_non_skills_via_classification(MockClient(), mock_tailored)
items = mock_tailored["tailored_resume"]["skills_selected"][0]["items"]
print(f"[OK] Remaining items: {items}")
assert "Wireshark" in items
assert "DWT Steganography" in items
assert "JWT Authentication" in items
assert "Stakeholder Communication" not in items, "Expected 'Stakeholder Communication' to be stripped"
assert "Requirements Gathering" not in items, "Expected 'Requirements Gathering' to be stripped"


# Test 2.22: Classification-based skill filtering with user selection priority
print("\nTEST 2.22: Classification-based skill filtering with user selection priority")

mock_tailored_user = {
    "tailored_resume": {
        "skills_selected": [
            {
                "category": "Skills",
                "items": ["Wireshark", "Stakeholder Communication", "Requirements Gathering"]
            }
        ]
    }
}

filter_non_skills_via_classification(
    MockClient(), 
    mock_tailored_user, 
    selected_keywords=["Stakeholder Communication"]
)
items_user = mock_tailored_user["tailored_resume"]["skills_selected"][0]["items"]
print(f"[OK] Remaining items with user selection: {items_user}")
assert "Wireshark" in items_user
assert "Stakeholder Communication" in items_user, "Expected 'Stakeholder Communication' to be preserved"
assert "Requirements Gathering" not in items_user, "Expected 'Requirements Gathering' to be stripped"




# Test 2.17: Fabrication check is non-blocking and generates unverified_skills correctly
print("\nTEST 2.17: Fabrication check is non-blocking and generates unverified_skills")
import gemini_service
import os
import json

original_generate_content = gemini_service.generate_content_with_fallback
original_api_key = os.environ.get("GEMINI_API_KEY")
os.environ["GEMINI_API_KEY"] = "mock_key"

class MockResponse:
    def __init__(self, text):
        self.text = text

# Mock output with fabricated skill "Express.js" and project technology "Kubernetes"
mock_output_fab = {
    "tailored_resume": {
        "summary": "Experienced engineer.",
        "skills_selected": [
            {"category": "Backend", "items": ["Python", "Java", "SQL", "C", "C++", "Express.js"]}
        ],
        "projects": [
            {
                "name": "SecureVault — Encrypted Secret Manager",
                "tech_stack": "Python, FastAPI, Kubernetes",
                "bullets": ["Wrote some code."]
            }
        ],
        "internships_selected": [],
        "education": [],
        "certifications": [],
        "section_order": ["skills", "projects"],
        "section_titles": {}
    },
    "entry_counts": {
        "master_projects_count": 1,
        "master_internships_count": 0,
        "master_education_count": 0
    },
    "dashboard": {}
}

def mock_gen_fab(*args, **kwargs):
    return MockResponse(json.dumps(mock_output_fab))

gemini_service.generate_content_with_fallback = mock_gen_fab

try:
    res = gemini_service.generate_tailored_resume(
        "Skills:\nPython\nJava\nSQL\nC\nC++\nPROJECTS:\n- SecureVault — Encrypted Secret Manager", 
        "Job requiring Python"
    )
    print("[OK] Generation completed successfully despite fabrication!")
    assert "unverified_skills" in res, "Expected 'unverified_skills' in the response!"
    unverified = res["unverified_skills"]
    print(f"Collected unverified items: {unverified}")
    
    # Check that Express.js and Kubernetes are unverified
    items_unverified = [x["item"] for x in unverified]
    assert "Express.js" in items_unverified, "Expected Express.js to be unverified"
    assert "Kubernetes" in items_unverified, "Expected Kubernetes to be unverified"
    
    # Confirm locations
    locations = {x["item"]: x["location"] for x in unverified}
    assert locations["Express.js"] == "skills category: Backend"
    assert locations["Kubernetes"] == "project: SecureVault — Encrypted Secret Manager"
except Exception as e:
    print(f"[FAIL] Test 2.17 failed: {e}")
    assert False, f"Expected successful generation, got error: {e}"
finally:
    gemini_service.generate_content_with_fallback = original_generate_content
    if original_api_key is not None:
        os.environ["GEMINI_API_KEY"] = original_api_key


# Test 2.18: Other hard-blocking checks (buzzwords, missing skills, cert count mismatch) still fail
print("\nTEST 2.18: Verify other hard-blocking checks still fail")
os.environ["GEMINI_API_KEY"] = "mock_key"

# Mock output with missing selected keyword
mock_output_missing_kw = {
    "tailored_resume": {
        "summary": "Experienced backend developer.",
        "skills_selected": [
            {"category": "Backend", "items": ["Python", "Java", "SQL", "C", "C++"]}
        ],
        "projects": [],
        "internships_selected": [],
        "education": [],
        "certifications": [],
        "section_order": ["skills"],
        "section_titles": {}
    },
    "entry_counts": {
        "master_projects_count": 0,
        "master_internships_count": 0,
        "master_education_count": 0
    },
    "dashboard": {}
}

def mock_gen_missing_kw(*args, **kwargs):
    return MockResponse(json.dumps(mock_output_missing_kw))

gemini_service.generate_content_with_fallback = mock_gen_missing_kw

try:
    res = gemini_service.generate_tailored_resume(
        "Skills:\nPython\nJava\nSQL\nC\nC++", 
        "Job requiring Python",
        ["Rust"]
    )
    assert False, "Expected missing keyword check to fail, but it succeeded!"
except ValueError as e:
    print(f"[OK] Missing keyword check failed as expected: {e}")
    assert "Missing user-selected keywords" in str(e)
finally:
    gemini_service.generate_content_with_fallback = original_generate_content
    if original_api_key is not None:
        os.environ["GEMINI_API_KEY"] = original_api_key









# ── Test 3: PDF rendering ─────────────────────────────────────────────────
print("\n" + "="*60)
print("TEST 3: PDF rendering (xhtml2pdf)")
print("="*60)

from pdf_service import render_pdf
pdf_path = render_pdf(MOCK_GEMINI_OUTPUT["tailored_resume"], MOCK_GEMINI_OUTPUT["tailored_resume"]["contact"])
assert pdf_path.exists(), f"PDF not created at {pdf_path}"
size_kb = pdf_path.stat().st_size / 1024
print(f"[OK] PDF created: {pdf_path.name}  ({size_kb:.1f} KB)")
assert size_kb > 3, f"PDF suspiciously small: {size_kb:.1f} KB"


# ── Test 4: Download path validation ─────────────────────────────────────
print("\n" + "="*60)
print("TEST 4: Download path validation")
print("="*60)

filename = pdf_path.name
assert re.match(r"^resume_\d{8}_\d{6}\.pdf$", filename), f"Bad filename: {filename}"
print(f"[OK] Filename format OK: {filename}")


# ── Test 5: Model Fallback Simulation ──────────────────────────────────────
print("\n" + "="*60)
print("TEST 5: Model Fallback Simulation")
print("="*60)

import gemini_service
from gemini_service import analyze_jd_match

# Temporarily prepend an invalid model name to simulate a 404/failure
original_models = list(gemini_service.GEMINI_MODELS)
gemini_service.GEMINI_MODELS = ["gemini-invalid-model-999"] + original_models
print(f"Configured models with fake primary model: {gemini_service.GEMINI_MODELS}")

try:
    # Run a simple match analysis (real API call)
    result = analyze_jd_match(
        master_resume="Alex Jordan\nPython, FastAPI, Docker",
        jd_text="Looking for a Python developer with FastAPI and Docker experience."
    )
    print(f"[OK] Fallback successful! Match Score: {result.get('overall_match_percentage')}%")
    assert "overall_match_percentage" in result, "Fallback result missing match percentage"
except Exception as exc:
    print(f"[FAIL] Fallback test failed: {exc}")
    raise
finally:
    # Restore original models
    gemini_service.GEMINI_MODELS = original_models


# ── Summary ───────────────────────────────────────────────────────────────
print("\n" + "="*60)
print("ALL TESTS PASSED [OK]")
print("="*60)
print(f"\nGenerated PDF: {pdf_path}")
print("The file upload -> text extraction -> safety check -> PDF pipeline -> validation loop works correctly.")
print("To test the full upload web interface, run uvicorn and open http://localhost:8000 in your browser.\n")
